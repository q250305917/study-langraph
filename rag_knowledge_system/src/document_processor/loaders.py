"""
文档加载器模块

提供多种格式的文档加载功能，支持PDF、Word、文本、Markdown等格式。
每个加载器负责将特定格式的文件转换为统一的Document对象。
"""

import os
import io
import tempfile
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Iterator
from dataclasses import dataclass
import chardet
import logging

# LangChain导入
from langchain.schema import Document
from langchain.document_loaders.base import BaseLoader

# 可选依赖导入
try:
    import PyPDF2
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup
    HAS_HTML = True
except ImportError:
    HAS_HTML = False

try:
    import markdownify
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

logger = logging.getLogger(__name__)


@dataclass
class LoaderConfig:
    """加载器配置"""
    encoding: str = "utf-8"
    extract_images: bool = False
    extract_tables: bool = False
    extract_metadata: bool = True
    chunk_size: int = 1024 * 1024  # 1MB
    timeout: int = 30


class BaseDocumentLoader(BaseLoader, ABC):
    """文档加载器基类"""
    
    def __init__(self, 
                 file_path: Union[str, Path],
                 config: Optional[LoaderConfig] = None):
        """
        初始化文档加载器
        
        Args:
            file_path: 文件路径
            config: 加载器配置
        """
        self.file_path = Path(file_path)
        self.config = config or LoaderConfig()
        
        # 验证文件存在
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {self.file_path}")
        
        # 验证文件格式
        if not self.is_supported_format():
            raise ValueError(f"不支持的文件格式: {self.file_path.suffix}")
    
    @abstractmethod
    def is_supported_format(self) -> bool:
        """检查是否支持当前文件格式"""
        pass
    
    @abstractmethod
    def _load_content(self) -> str:
        """加载文件内容（子类实现）"""
        pass
    
    def load(self) -> List[Document]:
        """
        加载文档
        
        Returns:
            Document列表
        """
        try:
            # 加载内容
            content = self._load_content()
            
            # 提取基本元数据
            metadata = self._extract_basic_metadata()
            
            # 创建Document对象
            document = Document(
                page_content=content,
                metadata=metadata
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"加载文档失败 {self.file_path}: {e}")
            raise
    
    def _extract_basic_metadata(self) -> Dict[str, Any]:
        """提取基本文件元数据"""
        stat = self.file_path.stat()
        
        metadata = {
            "source": str(self.file_path),
            "filename": self.file_path.name,
            "file_type": self.file_path.suffix.lower(),
            "file_size": stat.st_size,
            "created_time": stat.st_ctime,
            "modified_time": stat.st_mtime,
            "encoding": self.config.encoding
        }
        
        return metadata
    
    def lazy_load(self) -> Iterator[Document]:
        """惰性加载（适用于大文件）"""
        for doc in self.load():
            yield doc


class PDFLoader(BaseDocumentLoader):
    """PDF文档加载器"""
    
    def __init__(self,
                 file_path: Union[str, Path],
                 config: Optional[LoaderConfig] = None,
                 use_pdfplumber: bool = True):
        """
        初始化PDF加载器
        
        Args:
            file_path: PDF文件路径
            config: 加载器配置
            use_pdfplumber: 是否使用pdfplumber（更好的表格支持）
        """
        if not HAS_PDF:
            raise ImportError("需要安装PDF处理库: pip install PyPDF2 pdfplumber")
        
        super().__init__(file_path, config)
        self.use_pdfplumber = use_pdfplumber
    
    def is_supported_format(self) -> bool:
        """检查PDF格式支持"""
        return self.file_path.suffix.lower() == ".pdf"
    
    def _load_content(self) -> str:
        """加载PDF内容"""
        if self.use_pdfplumber and 'pdfplumber' in globals():
            return self._load_with_pdfplumber()
        else:
            return self._load_with_pypdf2()
    
    def _load_with_pdfplumber(self) -> str:
        """使用pdfplumber加载PDF"""
        content_parts = []
        
        with pdfplumber.open(self.file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    # 提取文本
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"=== 第{page_num}页 ===\n{text}")
                    
                    # 提取表格（如果配置启用）
                    if self.config.extract_tables:
                        tables = page.extract_tables()
                        for table_num, table in enumerate(tables, 1):
                            if table:
                                table_text = self._format_table(table)
                                content_parts.append(
                                    f"\n--- 表格{table_num} ---\n{table_text}"
                                )
                
                except Exception as e:
                    logger.warning(f"提取PDF第{page_num}页失败: {e}")
                    continue
        
        return "\n\n".join(content_parts)
    
    def _load_with_pypdf2(self) -> str:
        """使用PyPDF2加载PDF"""
        content_parts = []
        
        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"=== 第{page_num}页 ===\n{text}")
                
                except Exception as e:
                    logger.warning(f"提取PDF第{page_num}页失败: {e}")
                    continue
        
        return "\n\n".join(content_parts)
    
    def _format_table(self, table: List[List[str]]) -> str:
        """格式化表格为文本"""
        if not table:
            return ""
        
        # 计算每列的最大宽度
        col_widths = []
        for col_idx in range(len(table[0]) if table else 0):
            max_width = max(
                len(str(row[col_idx] or "")) 
                for row in table 
                if col_idx < len(row)
            )
            col_widths.append(max_width)
        
        # 格式化表格
        formatted_rows = []
        for row in table:
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                cell_str = str(cell or "")
                width = col_widths[col_idx] if col_idx < len(col_widths) else len(cell_str)
                formatted_cells.append(cell_str.ljust(width))
            formatted_rows.append(" | ".join(formatted_cells))
        
        return "\n".join(formatted_rows)
    
    def _extract_basic_metadata(self) -> Dict[str, Any]:
        """提取PDF元数据"""
        metadata = super()._extract_basic_metadata()
        
        try:
            if self.use_pdfplumber:
                with pdfplumber.open(self.file_path) as pdf:
                    metadata.update({
                        "total_pages": len(pdf.pages),
                        "pdf_version": getattr(pdf, 'pdf_version', None),
                        "metadata": pdf.metadata or {}
                    })
            else:
                with open(self.file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        "total_pages": len(pdf_reader.pages),
                        "metadata": pdf_reader.metadata or {}
                    })
        
        except Exception as e:
            logger.warning(f"提取PDF元数据失败: {e}")
        
        return metadata


class DocxLoader(BaseDocumentLoader):
    """Word文档加载器"""
    
    def __init__(self,
                 file_path: Union[str, Path],
                 config: Optional[LoaderConfig] = None):
        """初始化Word加载器"""
        if not HAS_DOCX:
            raise ImportError("需要安装Word处理库: pip install python-docx")
        
        super().__init__(file_path, config)
    
    def is_supported_format(self) -> bool:
        """检查Word格式支持"""
        return self.file_path.suffix.lower() in [".docx", ".doc"]
    
    def _load_content(self) -> str:
        """加载Word内容"""
        content_parts = []
        
        try:
            doc = DocxDocument(self.file_path)
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # 提取表格（如果配置启用）
            if self.config.extract_tables:
                for table_num, table in enumerate(doc.tables, 1):
                    table_text = self._format_docx_table(table)
                    if table_text:
                        content_parts.append(f"\n--- 表格{table_num} ---\n{table_text}")
        
        except Exception as e:
            logger.error(f"解析Word文档失败: {e}")
            raise
        
        return "\n".join(content_parts)
    
    def _format_docx_table(self, table) -> str:
        """格式化Word表格"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        
        return "\n".join(rows)
    
    def _extract_basic_metadata(self) -> Dict[str, Any]:
        """提取Word元数据"""
        metadata = super()._extract_basic_metadata()
        
        try:
            doc = DocxDocument(self.file_path)
            core_props = doc.core_properties
            
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": core_props.created,
                "modified": core_props.modified,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables) if self.config.extract_tables else 0
            })
        
        except Exception as e:
            logger.warning(f"提取Word元数据失败: {e}")
        
        return metadata


class TextLoader(BaseDocumentLoader):
    """纯文本文档加载器"""
    
    def is_supported_format(self) -> bool:
        """检查文本格式支持"""
        return self.file_path.suffix.lower() in [".txt", ".text", ".log"]
    
    def _load_content(self) -> str:
        """加载文本内容"""
        # 自动检测编码
        encoding = self._detect_encoding()
        
        try:
            with open(self.file_path, 'r', encoding=encoding) as file:
                return file.read()
        
        except UnicodeDecodeError:
            # 备用编码尝试
            for backup_encoding in ['utf-8', 'gbk', 'gb2312', 'latin1']:
                try:
                    with open(self.file_path, 'r', encoding=backup_encoding) as file:
                        logger.warning(f"使用备用编码 {backup_encoding} 加载文件")
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"无法确定文件编码: {self.file_path}")
    
    def _detect_encoding(self) -> str:
        """检测文件编码"""
        try:
            with open(self.file_path, 'rb') as file:
                raw_data = file.read(min(32768, self.config.chunk_size))  # 读取前32KB
                result = chardet.detect(raw_data)
                
                encoding = result.get('encoding', 'utf-8')
                confidence = result.get('confidence', 0)
                
                # 如果置信度太低，使用默认编码
                if confidence < 0.7:
                    encoding = self.config.encoding
                
                logger.debug(f"检测到编码: {encoding} (置信度: {confidence:.2f})")
                return encoding
        
        except Exception as e:
            logger.warning(f"编码检测失败: {e}")
            return self.config.encoding
    
    def _extract_basic_metadata(self) -> Dict[str, Any]:
        """提取文本元数据"""
        metadata = super()._extract_basic_metadata()
        
        try:
            content = self._load_content()
            lines = content.split('\n')
            
            metadata.update({
                "line_count": len(lines),
                "char_count": len(content),
                "word_count": len(content.split()),
                "detected_encoding": self._detect_encoding()
            })
        
        except Exception as e:
            logger.warning(f"提取文本元数据失败: {e}")
        
        return metadata


class MarkdownLoader(BaseDocumentLoader):
    """Markdown文档加载器"""
    
    def is_supported_format(self) -> bool:
        """检查Markdown格式支持"""
        return self.file_path.suffix.lower() in [".md", ".markdown", ".mdown", ".mkd"]
    
    def _load_content(self) -> str:
        """加载Markdown内容"""
        # 自动检测编码
        encoding = self._detect_encoding()
        
        with open(self.file_path, 'r', encoding=encoding) as file:
            content = file.read()
        
        # 可选择保持Markdown格式或转换为纯文本
        if hasattr(self.config, 'preserve_markdown') and self.config.preserve_markdown:
            return content
        else:
            # 转换为纯文本（保留结构）
            return self._markdown_to_text(content)
    
    def _markdown_to_text(self, content: str) -> str:
        """将Markdown转换为结构化文本"""
        lines = content.split('\n')
        processed_lines = []
        
        for line in lines:
            # 处理标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                processed_lines.append(f"{'='*level} {title} {'='*level}")
            
            # 处理列表
            elif line.strip().startswith(('*', '-', '+')):
                processed_lines.append(line)
            
            # 处理代码块
            elif line.strip().startswith('```'):
                processed_lines.append("--- 代码块 ---")
            
            # 普通内容
            else:
                processed_lines.append(line)
        
        return '\n'.join(processed_lines)
    
    def _detect_encoding(self) -> str:
        """检测Markdown文件编码"""
        try:
            with open(self.file_path, 'rb') as file:
                raw_data = file.read(8192)
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8')
        except:
            return 'utf-8'
    
    def _extract_basic_metadata(self) -> Dict[str, Any]:
        """提取Markdown元数据"""
        metadata = super()._extract_basic_metadata()
        
        try:
            content = self._load_content()
            
            # 提取标题信息
            lines = content.split('\n')
            headings = []
            for line in lines:
                if line.strip().startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    headings.append({"level": level, "title": title})
            
            metadata.update({
                "headings": headings,
                "heading_count": len(headings),
                "line_count": len(lines)
            })
        
        except Exception as e:
            logger.warning(f"提取Markdown元数据失败: {e}")
        
        return metadata


class HTMLLoader(BaseDocumentLoader):
    """HTML文档加载器"""
    
    def __init__(self,
                 file_path: Union[str, Path],
                 config: Optional[LoaderConfig] = None):
        """初始化HTML加载器"""
        if not HAS_HTML:
            raise ImportError("需要安装HTML处理库: pip install beautifulsoup4")
        
        super().__init__(file_path, config)
    
    def is_supported_format(self) -> bool:
        """检查HTML格式支持"""
        return self.file_path.suffix.lower() in [".html", ".htm"]
    
    def _load_content(self) -> str:
        """加载HTML内容"""
        encoding = self._detect_encoding()
        
        with open(self.file_path, 'r', encoding=encoding) as file:
            html_content = file.read()
        
        # 解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 移除script和style标签
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 提取文本
        text = soup.get_text()
        
        # 清理文本
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def _detect_encoding(self) -> str:
        """检测HTML编码"""
        try:
            with open(self.file_path, 'rb') as file:
                raw_data = file.read(8192)
                
                # 尝试从HTML meta标签中提取编码
                soup = BeautifulSoup(raw_data, 'html.parser')
                meta_charset = soup.find('meta', attrs={'charset': True})
                if meta_charset:
                    return meta_charset['charset']
                
                meta_content = soup.find('meta', attrs={'content': True})
                if meta_content and 'charset=' in meta_content.get('content', ''):
                    content = meta_content['content']
                    charset_pos = content.find('charset=')
                    if charset_pos != -1:
                        charset = content[charset_pos + 8:].split(';')[0]
                        return charset
                
                # 使用chardet检测
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8')
        
        except:
            return 'utf-8'
    
    def _extract_basic_metadata(self) -> Dict[str, Any]:
        """提取HTML元数据"""
        metadata = super()._extract_basic_metadata()
        
        try:
            with open(self.file_path, 'r', encoding=self._detect_encoding()) as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 提取基本信息
            title = soup.find('title')
            description = soup.find('meta', attrs={'name': 'description'})
            keywords = soup.find('meta', attrs={'name': 'keywords'})
            
            metadata.update({
                "title": title.get_text() if title else "",
                "description": description.get('content', '') if description else "",
                "keywords": keywords.get('content', '') if keywords else "",
                "lang": soup.find('html', attrs={'lang': True}).get('lang', '') if soup.find('html', attrs={'lang': True}) else ""
            })
        
        except Exception as e:
            logger.warning(f"提取HTML元数据失败: {e}")
        
        return metadata


# 工厂函数
def get_loader_for_file(file_path: Union[str, Path], 
                       config: Optional[LoaderConfig] = None) -> BaseDocumentLoader:
    """
    根据文件类型自动选择合适的加载器
    
    Args:
        file_path: 文件路径
        config: 加载器配置
        
    Returns:
        对应的文档加载器实例
        
    Raises:
        ValueError: 不支持的文件格式
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    loader_map = {
        '.pdf': PDFLoader,
        '.docx': DocxLoader,
        '.doc': DocxLoader,
        '.txt': TextLoader,
        '.text': TextLoader,
        '.log': TextLoader,
        '.md': MarkdownLoader,
        '.markdown': MarkdownLoader,
        '.mdown': MarkdownLoader,
        '.mkd': MarkdownLoader,
        '.html': HTMLLoader,
        '.htm': HTMLLoader,
    }
    
    if suffix not in loader_map:
        raise ValueError(f"不支持的文件格式: {suffix}")
    
    loader_class = loader_map[suffix]
    return loader_class(file_path, config)


def load_document(file_path: Union[str, Path], 
                 config: Optional[LoaderConfig] = None) -> List[Document]:
    """
    加载单个文档的便捷函数
    
    Args:
        file_path: 文件路径
        config: 加载器配置
        
    Returns:
        Document列表
    """
    loader = get_loader_for_file(file_path, config)
    return loader.load()


def load_documents_from_directory(directory: Union[str, Path],
                                 config: Optional[LoaderConfig] = None,
                                 recursive: bool = True,
                                 supported_extensions: Optional[List[str]] = None) -> List[Document]:
    """
    从目录批量加载文档
    
    Args:
        directory: 目录路径
        config: 加载器配置
        recursive: 是否递归搜索子目录
        supported_extensions: 支持的文件扩展名列表
        
    Returns:
        Document列表
    """
    directory = Path(directory)
    if not directory.exists() or not directory.is_dir():
        raise ValueError(f"目录不存在或不是有效目录: {directory}")
    
    if supported_extensions is None:
        supported_extensions = ['.pdf', '.docx', '.doc', '.txt', '.md', '.html', '.htm']
    
    documents = []
    
    # 搜索文件
    pattern = "**/*" if recursive else "*"
    
    for file_path in directory.glob(pattern):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            try:
                docs = load_document(file_path, config)
                documents.extend(docs)
                logger.info(f"成功加载文档: {file_path}")
            
            except Exception as e:
                logger.error(f"加载文档失败 {file_path}: {e}")
                continue
    
    logger.info(f"从目录 {directory} 共加载 {len(documents)} 个文档")
    return documents